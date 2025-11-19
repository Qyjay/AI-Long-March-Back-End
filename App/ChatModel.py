import os
import json
import asyncio
import tempfile
from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import JSONLoader
import dashscope  
# LangChain 新旧版本导入兼容
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.retrievers import BaseRetriever
from langchain_core.documents import Document as LcDocument
from typing import List, Any, Dict
from asyncio import Queue
from .core.config import get_settings
from typing import Optional

# 设置环境变量解决OpenMP冲突
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'

def get_llm() -> ChatOpenAI:
    """创建并返回聊天模型实例，支持 OpenAI 与 DashScope 兼容模式"""
    settings = get_settings()
    api_key = settings.OPENAI_API_KEY or settings.DASHSCOPE_API_KEY
    if not api_key:
        raise ValueError("缺少 OPENAI_API_KEY 或 DASHSCOPE_API_KEY")
    global _llm
    try:
        return _llm
    except NameError:
        use_dashscope = bool(settings.DASHSCOPE_API_KEY)
        model_name = settings.LLM_MODEL
        base_url = settings.OPENAI_BASE_URL if settings.OPENAI_BASE_URL else (settings.OPENAI_BASE_URL if use_dashscope else None)
        _llm = ChatOpenAI(
            model=model_name,
            openai_api_key=api_key,
            base_url=base_url,
            streaming=True,
            temperature=0,
        )
        return _llm

class DashScopeEmbeddings:
    """DashScope嵌入实现，兼容 LangChain 接口"""
    
    def __init__(self, model: str = "text-embedding-v4", api_key: str = None):
        """构造并配置嵌入模型与密钥"""
        settings = get_settings()
        self.model = model or settings.EMBEDDING_MODEL
        settings = get_settings()
        self.api_key = api_key or settings.DASHSCOPE_API_KEY
        dashscope.api_key = self.api_key
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """生成文档嵌入向量"""
        embeddings = []
        for text in texts:
            resp = dashscope.TextEmbedding.call(model=self.model, input=text)
            if hasattr(resp, "status_code") and resp.status_code == 200:
                embedding = resp.output['embeddings'][0]['embedding']
                embeddings.append(embedding)
            else:
                embeddings.append([0.0] * 1536)
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        """生成查询嵌入向量"""
        resp = dashscope.TextEmbedding.call(model=self.model, input=text)
        if hasattr(resp, "status_code") and resp.status_code == 200:
            return resp.output['embeddings'][0]['embedding']
        return [0.0] * 1536

class SafeEmbeddings:
    """安全嵌入包装器：在底层嵌入调用失败时回退为零向量。

    提供与 LangChain 兼容的 `embed_documents` 与 `embed_query` 接口，确保
    RAG 初始化与检索在 API 密钥无效或网络异常时依然可用。
    """
    def __init__(self, base, dim: int = 1536):
        """构造包装器并指定维度"""
        self._base = base
        self._dim = dim

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量生成文档向量，失败时回退为零向量"""
        try:
            return self._base.embed_documents(texts)
        except Exception:
            return [[0.0] * self._dim for _ in texts]

    def embed_query(self, text: str) -> List[float]:
        """生成查询向量，失败时回退为零向量"""
        try:
            return self._base.embed_query(text)
        except Exception:
            return [0.0] * self._dim

def get_embeddings():
    """返回可用的嵌入实现（DashScope 或 OpenAI），并使用安全包装器"""
    settings = get_settings()
    if settings.DASHSCOPE_API_KEY:
        return SafeEmbeddings(DashScopeEmbeddings(model=settings.EMBEDDING_MODEL))
    if settings.OPENAI_API_KEY:
        base = OpenAIEmbeddings(model="text-embedding-3-small", api_key=settings.OPENAI_API_KEY, base_url=settings.OPENAI_BASE_URL)
        # text-embedding-3-small 维度 1536
        return SafeEmbeddings(base, dim=1536)
    raise ValueError("缺少嵌入API密钥配置")
        
class StreamRetriever(BaseRetriever):
    """支持流式输出的检索器

    直接使用向量检索以规避不同版本 FAISS 对 `embedding_function` 的调用差异。
    """
    
    def __init__(self, vectorstore: FAISS, embeddings: Any, k: int = 3):
        """构造检索器，保存向量库与嵌入实现"""
        super().__init__()
        self._vectorstore = vectorstore
        self._embeddings = embeddings
        self._k = k
    
    def _get_relevant_documents(self, query: str) -> List[LcDocument]:
        """使用查询向量进行相似度检索"""
        try:
            qv = self._embeddings.embed_query(query)
            return self._vectorstore.similarity_search_by_vector(qv, k=self._k)
        except Exception:
            # 回退至文本检索
            return self._vectorstore.similarity_search(query, k=self._k)
    
    async def _aget_relevant_documents(self, query: str) -> List[LcDocument]:
        """异步版本"""
        return self._get_relevant_documents(query)

class RAGSystem:
    def __init__(self, docx_file_path: str, persist_directory: str = "./faiss_index"):
        self.docx_file_path = docx_file_path
        self.persist_directory = persist_directory
        self.vectorstore = None
        self.retriever = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
        )
        self._initialize_system()
    
    def _initialize_system(self):
        """初始化RAG系统"""
        try:
            settings = get_settings()
            self.embeddings = get_embeddings()
            
            # 尝试加载现有的向量数据库
            if os.path.exists(self.persist_directory):
                print("加载现有的FAISS索引...")
                # 为兼容旧版 FAISS 的 embedding_function 签名，传入可调用函数
                self.vectorstore = FAISS.load_local(
                    self.persist_directory,
                    lambda q: self.embeddings.embed_query(q),
                    allow_dangerous_deserialization=True,
                )
                print("FAISS索引加载成功")
            else:
                print("创建新的FAISS索引...")
                self._create_vectorstore()
            
            # 初始化检索器
            self.retriever = StreamRetriever(self.vectorstore, self.embeddings, k=3)
            
        except Exception as e:
            print(f"初始化RAG系统失败: {e}")
            # 如果加载失败，重新创建
            self._create_vectorstore()
            self.retriever = StreamRetriever(self.vectorstore, self.embeddings, k=3)
    
    def _create_vectorstore(self):
        """创建向量数据库"""
        try:
            # 加载DOCX文档
            print(f"加载DOCX文件: {self.docx_file_path}")
            
            # 检查文件是否存在
            if not os.path.exists(self.docx_file_path):
                print(f"DOCX文件不存在: {self.docx_file_path}")
                # 创建示例数据文件
                self._create_sample_data()
            
            # 使用python-docx加载文档
            from docx import Document as DocxDocument
            doc = DocxDocument(self.docx_file_path)
            
            # 提取文本内容
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # 生成元数据
            metadata = {
                "source": self.docx_file_path,
                "title": os.path.basename(self.docx_file_path)
            }
            
            # 创建文档对象
            documents = [LcDocument(page_content=text, metadata=metadata)]
            print(f"原始文档数量: {len(documents)}")
            
            if documents and documents[0].page_content.strip():
                # 分割文档
                splits = self.text_splitter.split_documents(documents)
                print(f"分割后文档数量: {len(splits)}")
                
                # 创建向量数据库（使用嵌入对象以支持批量嵌入）
                self.vectorstore = FAISS.from_documents(
                    splits,
                    self.embeddings,
                )
                
                # 保存向量数据库
                self.vectorstore.save_local(self.persist_directory)
                print(f"FAISS索引创建成功，保存到: {self.persist_directory}")
            else:
                # 创建包含默认文档的向量数据库
                print("没有加载到文档，创建默认文档...")
                self._create_default_vectorstore()
                
        except Exception as e:
            print(f"创建向量数据库失败: {e}")
            # 创建包含默认文档的向量数据库作为备用
            self._create_default_vectorstore()
    
    def _create_sample_data(self):
        """在缺失 DOCX 时创建一个最小示例 DOCX 文件。"""
        from docx import Document
        doc = Document()
        doc.add_heading("示例文档", level=1)
        doc.add_paragraph("这是一个用于初始化 RAG 索引的示例文档。")
        doc.save(self.docx_file_path)
        print(f"已创建示例 DOCX 文件: {self.docx_file_path}")
    
    def _create_default_vectorstore(self):
        """创建默认的向量数据库"""
        default_docs = [LcDocument(
            page_content="这是一个默认的文档，请确保您的DOCX文件存在且包含内容。",
            metadata={"source": "default", "category": "system"}
        )]
        self.vectorstore = FAISS.from_documents(
            default_docs,
            self.embeddings,
        )
        self.vectorstore.save_local(self.persist_directory)
        print("已创建默认向量数据库")
        
    def search_documents(self, query: str, k: int = 3) -> List[Dict]:
        """检索相关文档"""
        try:
            if self.retriever is None or self.vectorstore is None:
                return []
            
            # 使用新的 invoke 方法替代弃用的 get_relevant_documents
            docs = self.retriever.invoke(query)
            results = []
            
            for i, doc in enumerate(docs):
                results.append({
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'score': 1.0 - (i * 0.1)  # 简单的分数估算
                })
            
            return results
            
        except Exception as e:
            print(f"文档检索失败: {e}")
            return []
    



    def build_explanation_prompt(self, route: str, choice: str, real_history_choice: str, 
                           history_background: str, tactical_logic: str) -> str:
        """构建讲解员风格的提示词"""
        
        prompt_template = """角色：你是"AI陪伴下的重走长征路"网页中的智能讲解员小精灵，语气亲切、历史严谨，擅长用浅白语言解析历史逻辑。
    任务：用户在{route}节点做出了{choice}的选择，请结合以下历史事实，解析该选择与历史真实决策的差异，并解释历史决策的合理性：
    1. 历史真实决策：{real_history_choice}
    2. 历史背景：{history_background}
    3. 战术逻辑：{tactical_logic}

    输出要求：
    - 先肯定用户选择的思考（如"你考虑到保护中央纵队的重要性，这个思路很关键！"）
    - 再客观指出差异（如"不过历史上红军并没有选择正面突破，而是采用了分兵佯攻的战术"）
    - 最后解释历史决策的原因（结合上述历史背景、战术逻辑，避免专业术语）
    - 结尾关联长征精神（如"这种灵活调整的战术，正是长征中'实事求是'精神的体现"）

    请根据以上要求提供回答："""
        
        return prompt_template.format(
            route=route,
            choice=choice,
            real_history_choice=real_history_choice,
            history_background=history_background,
            tactical_logic=tactical_logic
        )
    def build_poem_prompt(self, route: str, poem: str, creation_background: str, 
                     poem_analysis: str, spirit: str, meanings: str) -> str:
        """构建诗词赏析风格的提示词"""
        
        prompt_template = """角色：你是红色诗词赏析专家，需兼顾文学性与历史背景，语言生动易懂，适合青少年用户。
    任务：为{route}节点的{poem}提供赏析，包含以下维度：
    1. 创作背景：{creation_background}
    2. 词句解析：{poem_analysis}
    3. 情感与精神：{spirit}
    4. 当代意义：{meanings}

    输出要求：
    - 分点但不用列表，用"创作背景""词句意思""精神内涵""现在的意义"作为引导词
    - 每部分不超过200字，避免学术化表达（如不说"意象"，说"用'铁'来形容关隘很难攻克"）

    请根据以上要求提供赏析："""
        
        return prompt_template.format(
            route=route,
            poem=poem,
            creation_background=creation_background,
            poem_analysis=poem_analysis,
            spirit=spirit,
            meanings=meanings
        )
    def build_qa_prompt(self, question: str, context_docs: List[Dict]) -> str:
        """构建问答风格的提示词"""
        
        if context_docs:
            context_text = "\n\n".join([
                f"【相关文档 {i+1}】\n{doc['content']}\n来源：{doc['metadata'].get('source', '未知')}" 
                for i, doc in enumerate(context_docs)
            ])
        else:
            context_text = "暂无相关文档"
        
        prompt_template = """角色：你是长征历史知识问答专家，需基于提供的RAG素材（标注Source）回答用户问题，确保信息准确、来源可查，语气友好。
    任务：用户提问{question}，请结合以下RAG素材回答：
    {context}

    输出要求：
    1. 先直接回应问题（如"是真的，但和我们平时理解的'吃皮带'不太一样"）
    2. 详细解释（结合RAG素材，如"红军吃的是皮带的内层纤维，需要先去掉表层皮革，煮软后才能食用"）
    3. 补充背景（如"当时青稞面吃完了，野菜也难找，皮带是无奈之下的选择，体现了红军的艰难"）
    4. 标注信息来源（如"参考《红军过草地纪实》中对物资短缺的记载"）
    5. 避免猜测：若素材未覆盖，需说明"目前没有明确历史记载，但根据当时情况推测……"

    请根据以上要求提供回答："""
        
        return prompt_template.format(
            question=question,
            context=context_text
        )
def _init_rag() -> RAGSystem:
    """初始化并返回 RAG 系统实例"""
    settings = get_settings()
    return RAGSystem(settings.RAG_DOCX_PATH, settings.FAISS_DIR)

rag_system = _init_rag()

class StreamProcessor:
    """处理流式输出的类"""
    
    def __init__(self):
        self.content_queue = Queue()
    
    async def process_stream(self, completion, context_docs):
        """处理流式响应"""
        try:
            # 发送开始信号和检索上下文
            await self.content_queue.put({
                "type": "start",
                "context_docs": context_docs
            })
            
            full_content = ""
            usage_info = None
            
            # 逐块处理流式响应
            for chunk in completion:
                # 兼容 LangChain ChatOpenAI（AIMessageChunk）
                content_attr = getattr(chunk, "content", None)
                if content_attr is not None:
                    content = content_attr
                    full_content += content
                    await self.content_queue.put({
                        "type": "content",
                        "content": content,
                        "full_content": full_content
                    })
                    await asyncio.sleep(0.01)
                    continue

                # 兼容 OpenAI 原始流（choices[0].delta.content）
                try:
                    if chunk.choices and chunk.choices[0].delta.content is not None:
                        content = chunk.choices[0].delta.content
                        full_content += content
                        await self.content_queue.put({
                            "type": "content",
                            "content": content,
                            "full_content": full_content
                        })
                        await asyncio.sleep(0.01)
                        continue
                except Exception:
                    pass
            
            # 发送结束信号
            await self.content_queue.put({
                "type": "complete",
                "full_content": full_content,
                "usage": usage_info,
                "context_docs": context_docs
            })
            
        except Exception as e:
            await self.content_queue.put({
                "type": "error",
                "message": str(e)
            })
    
    async def generate_responses(self):
        """生成响应数据"""
        while True:
            try:
                data = await self.content_queue.get()
                if data is None:  # 结束信号
                    break
                
                yield "data: " + json.dumps(data, ensure_ascii=False) + "\n\n"
                
            except Exception as e:
                error_data = {"type": "error", "message": str(e)}
                yield "data: " + json.dumps(error_data, ensure_ascii=False) + "\n\n"

    async def process_fallback(self, context_docs, text: str, chunk_size: int = 120):
        """在LLM不可用时的降级处理：将给定文本切块并按SSE事件推送

        参数:
        - context_docs: RAG检索到的文档列表，用于在开始事件中附带
        - text: 需要流式输出的降级文本内容
        - chunk_size: 每次输出的字符块大小
        """
        try:
            # 发送开始事件
            await self.content_queue.put({
                "type": "start",
                "context_docs": context_docs
            })
            full = ""
            # 切块输出内容
            for i in range(0, len(text), chunk_size):
                part = text[i:i+chunk_size]
                full += part
                await self.content_queue.put({
                    "type": "content",
                    "content": part,
                    "full_content": full
                })
                await asyncio.sleep(0.02)
            # 完成事件
            await self.content_queue.put({
                "type": "complete",
                "full_content": full,
                "usage": None,
                "context_docs": context_docs
            })
        except Exception as e:
            await self.content_queue.put({
                "type": "error",
                "message": str(e)
            })
